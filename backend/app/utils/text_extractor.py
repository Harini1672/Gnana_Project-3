import logging
import io
import re
import pdfplumber
import PyPDF2
import docx
import pandas as pd

logger = logging.getLogger("app.text_extractor")


def _clean_pdf_text(text: str) -> str:
    """Light post-processing on PDF text: fix hyphenated line-breaks
    and collapse excessive blank lines.  Does NOT attempt to fix
    mid-word spaces — pdfplumber handles that at extraction time.
    """
    # Rejoin hyphenated line-breaks: "learn-\ning" → "learning"
    text = re.sub(r"-\s*\n\s*", "", text)
    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber (much cleaner than PyPDF2).

    pdfplumber uses pdfminer under the hood and produces proper word-level
    layout, avoiding the mid-word space artefacts PyPDF2 creates.
    Falls back to PyPDF2 only if pdfplumber fails entirely.
    """
    try:
        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=2, y_tolerance=3)
                if text:
                    pages.append(text)
        raw = "\n".join(pages)
        if not raw.strip():
            raise ValueError("pdfplumber returned empty text.")
        return _clean_pdf_text(raw)
    except Exception as primary_exc:
        logger.warning(
            "pdfplumber extraction failed (%s) — falling back to PyPDF2.", primary_exc
        )
        try:
            import PyPDF2
            pdf_file = io.BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file)
            pages2 = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages2.append(t)
            raw2 = "\n".join(pages2)
            return _clean_pdf_text(raw2)
        except Exception as fallback_exc:
            logger.error("Both PDF extractors failed: %s", fallback_exc)
            raise ValueError(
                f"Could not extract text from PDF: {fallback_exc}"
            ) from fallback_exc

def extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes and clean PyPDF2 artefacts."""
    text_content = []
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
        raw = "\n".join(text_content)
        return _clean_pdf_text(raw)
    except Exception as e:
        logger.error("Error extracting PDF: %s", e)
        raise ValueError(f"Could not extract text from PDF file: {e}")

def extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes, including paragraph structures and tables."""
    text_content = []
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
                
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells_text:
                    text_content.append(" | ".join(cells_text))
                    
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        raise ValueError(f"Could not extract text from DOCX file: {e}")

def extract_csv(file_bytes: bytes) -> str:
    """Extract text from CSV file bytes, representing each row as structured key-value sentences."""
    try:
        csv_file = io.BytesIO(file_bytes)
        df = pd.read_csv(csv_file)
        
        lines = []
        for idx, row in df.iterrows():
            row_items = []
            for col_name, val in row.items():
                # Format float numbers nicely
                if isinstance(val, float):
                    row_items.append(f"{col_name}: {val:.2f}")
                else:
                    row_items.append(f"{col_name}: {str(val).strip()}")
            lines.append(", ".join(row_items))
            
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"Error extracting CSV: {e}")
        raise ValueError(f"Could not extract text from CSV file: {e}")

def extract_txt(file_bytes: bytes) -> str:
    """Extract text from TXT file bytes using utf-8 decoding (fallback to latin-1)."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception as e:
            logger.error(f"Error decoding TXT: {e}")
            raise ValueError("Unsupported text encoding. Please upload files in UTF-8 or standard Latin-1.")

def extract_text(file_bytes: bytes, file_name: str) -> str:
    """Detect file extension and extract plain text from file bytes."""
    ext = file_name.split(".")[-1].lower()
    
    if ext == "pdf":
        return extract_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return extract_docx(file_bytes)
    elif ext == "csv":
        return extract_csv(file_bytes)
    elif ext in ["txt", "md"]:
        return extract_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension: .{ext}. Supported formats are PDF, DOCX, CSV, TXT.")
